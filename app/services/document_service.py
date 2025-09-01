import os
import aiofiles
from pathlib import Path
from typing import List, Optional
import uuid
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from io import BytesIO

from sqlalchemy.orm import Session
from app.models.documents import Document, DocumentChunk
from app.services.embedding_service import embedding_service
from app.services.vector_service import vector_service

class DocumentService:
    def __init__(self):
        # Create documents directory
        self.upload_dir = Path("data/documents")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Allowed file extensions
        self.allowed_extensions = {".pdf", ".docx", ".txt", ".md"}
        
        # Text chunking settings
        self.chunk_size = 1000  # characters per chunk
        self.chunk_overlap = 200  # overlap between chunks
    
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """
        Save uploaded file to disk
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Path to saved file
        """
        # Generate unique filename
        file_extension = Path(filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = self.upload_dir / unique_filename
        
        # Save file asynchronously
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(file_content)
        
        return str(file_path)
    
    def extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from various file types
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == ".pdf":
            return self._extract_from_pdf(file_path)
        elif file_extension == ".docx":
            return self._extract_from_docx(file_path)
        elif file_extension in [".txt", ".md"]:
            return self._extract_from_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        with open(file_path, "rb") as file:
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks for better search
        
        Args:
            text: Input text to chunk
            
        Returns:
            List of text chunks
        """
        chunks = []
        text_length = len(text)
        
        start = 0
        while start < text_length:
            # Calculate chunk end
            end = start + self.chunk_size
            
            # If not at the end, try to break at word boundary
            if end < text_length:
                # Find last space within chunk to avoid breaking words
                last_space = text.rfind(" ", start, end)
                if last_space > start:
                    end = last_space
            
            # Extract chunk
            chunk = text[start:end].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start <= 0:
                start = end
        
        return chunks
    
    async def process_document(
        self, 
        db: Session,
        file_content: bytes,
        filename: str,
        title: str,
        department: str,
        content_type: str,
        uploaded_by: str
    ) -> Document:
        """
        Complete document processing pipeline
        
        Args:
            db: Database session
            file_content: File content as bytes
            filename: Original filename
            title: Document title
            department: Department that owns the document
            content_type: Type of document (policy, manual, etc.)
            uploaded_by: User who uploaded the document
            
        Returns:
            Created Document object
        """
        # 1. Save file to disk
        file_path = await self.save_uploaded_file(file_content, filename)
        
        # 2. Extract text content
        text_content = self.extract_text_from_file(file_path)
        
        # 3. Create document record in database
        document = Document(
            title=title,
            file_path=file_path,
            department=department,
            content_type=content_type,
            file_size=len(file_content),
            original_filename=filename,
            uploaded_by=uploaded_by
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # 4. Chunk the text
        chunks = self.chunk_text(text_content)
        
        # 5. Generate embeddings and store chunks
        for i, chunk_text in enumerate(chunks):
            # Generate embedding for this chunk
            embedding = embedding_service.generate_embedding(chunk_text)
            
            # Store in vector database
            chunk_id = vector_service.add_document_chunk(
                chunk_text=chunk_text,
                embedding=embedding,
                metadata={
                    "document_id": document.id,
                    "chunk_index": i,
                    "title": title,
                    "department": department,
                    "content_type": content_type
                }
            )
            
            # Store chunk metadata in SQL database
            document_chunk = DocumentChunk(
                document_id=document.id,
                content=chunk_text,
                chunk_index=i,
                start_char=i * (self.chunk_size - self.chunk_overlap),
                end_char=(i + 1) * (self.chunk_size - self.chunk_overlap),
                vector_id=chunk_id
            )
            
            db.add(document_chunk)
        
        db.commit()
        return document

# Global instance
document_service = DocumentService()